import os
import docx
import nltk
from pprint import pprint
import matplotlib.pyplot as plt
import numpy as np
import contractions
from contractions import contractions_dict
import re
from nltk.stem import LancasterStemmer
import unicodedata
from nltk.stem import WordNetLemmatizer
import spacy
from spacy.lang.en import English
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.feature_extraction.text import TfidfTransformer
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import pandas as pd
from IPython.display import display
from gensim.models import word2vec
from scipy.cluster.hierarchy import dendrogram, linkage, fcluster
from sklearn.decomposition import LatentDirichletAllocation



# open file docx
def read_docx(text, paragraph=False):
    doc = docx.Document(text)
    text1 = ""

    if not paragraph:
        return doc.paragraphs
    else:
        for i in range(len(doc.paragraphs)):
            if doc.paragraphs[i].text == "":
                continue
            text1 = text1 + " " + doc.paragraphs[i].text
        return text1


# word and sentence tokenization
def tokenize_text(text):
    sentences = nltk.sent_tokenize(text=text)
    word_tokens = [nltk.word_tokenize(sentence) for sentence in sentences]
    return word_tokens


# expands shortend word like I'll -> I will
def expand_contractions(text, contraction_mapping=contractions_dict):
    contraction_pattern = re.compile('({})'.format('|'.join(contraction_mapping.keys())), flags=re.IGNORECASE | re.DOTALL)

    def expand_match(contraction):
        match = contraction.group(0)
        first_char = match[0]
        expand_contractions = contraction_mapping.get(match) if contraction_mapping.get(match) \
            else contraction_mapping.get(match.lower())
        expanded_contraction = first_char + expand_contractions[1:]
        return expanded_contraction

    expanded_text = contraction_pattern.sub(expand_match, text)
    expanded_text = re.sub("'", "", expanded_text)
    return expanded_text


def expand_contractions2(text):
    expanded_words = []
    for word in text.split():
        expanded_words.append(contractions.fix(word))
    expanded_text = ' '.join(expanded_words)
    return expanded_text


def remove_accented_chars(text):
    text = unicodedata.normalize('NFKD', text).encode('ascii', 'ignore').decode('utf-8', 'ignore')
    return text


# sentence tokenization
def sentence_tokenization(text):
    default_st = nltk.sent_tokenize
    text_sentence = default_st(text=text)
    print('Total sentences in sample_text:', len(text_sentence))
    # pipreturn np.array(text_sentence)
    return text_sentence


def word_tokenization(text):  # word tokenization
    default_wt = nltk.word_tokenize
    text_words = default_wt(text=text)
    words_array = np.array(text_words)
    print(len(words_array))
    return words_array


# Remove special characters
def remove_special_char(text, remove_digits=False):
    # pattern = r'[^a-zA-z0-9\s]' if not remove_digits else r'[^a-zA-z\s]'
    if not remove_digits:
        pattern = r'[^a-zA-z0-9\s]'
    else:
        pattern = r'[^a-zA-z\s]'
    text = re.sub(pattern, '', text)
    return text


def stopwords_remove(text, is_lowr_case=False):
    tokenizer = nltk.ToktokTokenizer()
    stopwords_list = nltk.corpus.stopwords.words('english')
    tokens = tokenizer.tokenize(text)
    tokens = [token.strip() for token in tokens]
    if is_lowr_case:
        filtered_tokens = [token for token in tokens if token not in stopwords_list]
    else:
        filtered_tokens = [token for token in tokens if token.lower() not in stopwords_list]
    filtered_text = ' '.join(filtered_tokens)
    return filtered_text


def lemmatize(text):
    nlp = spacy.load("en_core_web_lg")
    # lemmatizer = nlp.get_pipe("lemmatizer")
    doc = nlp(text)
    # docs = list(nlp.pipe(text))
    x = [token.lemma_ for token in doc]

    return x

# zrychlit lemmatizaci
# normalizace textu mi rozdělí text na věty a provede na nich upravy


def normalizace_textu(text, contraction_expansion=True, accented_char_removal=True,
                      text_lemmatization=True, special_char_removal=True, stopword_removal=True,
                      remove_digits=True, text_lower_case=True):

    normalized_text = []
    for doc in text:
        if accented_char_removal:
            doc = remove_accented_chars(doc)
        if contraction_expansion:
            doc = expand_contractions2(doc)
        if text_lower_case:
            doc = doc.lower()
        if special_char_removal:
            doc = remove_special_char(doc)
        if stopword_removal:
            doc = stopwords_remove(doc)
        if remove_digits:
            doc = remove_special_char(doc, remove_digits=True)
        doc = re.sub(r'[\r|\n|\r\n]+', ' ', doc)
        # if len(doc.split()) <= 1:
       # remove doc from text array
       # if text_lemmatization:
        #    doc = lemmatize(doc)

        normalized_text.append(doc)

    return normalized_text


# ----------------------------------BOW TF-IDF----------------------------------------------------------------------


def BOW(text, number_of_most_used_words, if_print=True):

    text = sentence_tokenization(text)

    for i in range(len(text)):
        print(text[i])
    print(type(text))
    y = normalizace_textu(text)
    print(" ")
    print(type(y))
    print("")
    y = [string for string in y if string != '']

    for j in range(len(y)):
        if if_print:
            print(y[j])

    y = lemmatize(" ".join(y))
    if if_print:
        print(len(y))
    y = [string for string in y if string != ' ']
    if if_print:
        print(len(y))
        print(y)

    cv = CountVectorizer(min_df=0.,max_df=1.)
    cv_matrix =cv.fit_transform(y)

    m = np.sum(cv_matrix, axis=0)
    if if_print:
        print(type(cv_matrix))
    names = cv.get_feature_names_out()
    display(pd.DataFrame(m, columns=names))

    l = []
    for i in range(0, m.size):
        l.append(m[0, i])
    if if_print:
        print(l)

    c = []
    for i in range(0, m.size):
        c.append(names[i])
    if if_print:
        print(c)

    n = []
    for j in range(0, m.size):
        s = str(l[j]) + " " + str(c[j])
        n.append(s.split())

    for element in n:
        element[0] = int(element[0])
    n.sort(reverse=True)
    if if_print:
        print(n[:number_of_most_used_words])
    return cv_matrix, cv


def multiple_text(text1, text2):
    full_text = read_docx(text1, paragraph=True) + read_docx(text2, paragraph=True)
    return full_text


def tf_idf_t(BOW_MATRIX, Count_Vectorizer):
    transform = TfidfTransformer(norm='l2', use_idf=True)
    transform_matrix = transform.fit_transform(BOW_MATRIX)
    m = np.sum(transform_matrix, axis=0)
    transform_matrix = transform_matrix.toarray()
    names = Count_Vectorizer.get_feature_names_out()
    DF1 = pd.DataFrame(np.round(transform_matrix, 2), columns=names)
    display(DF1)
    DF1.to_csv("TFIDF.csv", index=False)

    l = []
    for i in range(0, m.size):
        l.append(m[0, i])
    print("this is l",l)

    c = []
    for i in range(0, m.size):
        c.append(names[i])
    print("this is c", c)

    n = []
    for j in range(0, m.size):
        s = str(l[j]) + " " + str(c[j])
        n.append(s.split())

    for element in n:
        element[0] = float(element[0])
    n.sort(reverse=True)
    print(n[:10])


def tf_idf_v(text, rm_stop_words=False, not_tokenized=True):
    if not_tokenized:
        text = sentence_tokenization(text)
    if rm_stop_words:
        y = normalizace_textu(text)
        print(" ")
        print(type(y))
        print("")
        y = [string for string in y if string != '']

        for j in range(len(y)):
            print(y[j])

        y = lemmatize(" ".join(y))
        print(len(y))
        print("after lemma",y)
        #y = [string for string in y if string != ' ']
        print(y)
        return 0
    else:
        tfidf_vectorizer = TfidfVectorizer(min_df=0., max_df=1., norm="l2", use_idf=True, smooth_idf=True)
        tfidf_vectorizer_matrix = tfidf_vectorizer.fit_transform(text).toarray()
        vocab = tfidf_vectorizer.get_feature_names_out()
        DF1 = pd.DataFrame(np.round(tfidf_vectorizer_matrix, 2), columns=vocab)
        display(pd.DataFrame(np.round(tfidf_vectorizer_matrix, 2), columns=vocab))
        DF1.to_csv("TFIDF.csv", index=False)

    return tfidf_vectorizer_matrix, DF1, vocab


def tf_idf_cosine_similarity_two_doc(a, b):
    shape = np.maximum(a.shape, b.shape)

    a1 = np.zeros(shape)
    b1 = np.zeros(shape)
    a1[:a.shape[0], :a.shape[1]] = a
    b1[:b.shape[0], :b.shape[1]] = b

    cosine_s_matrix = cosine_similarity(a1, b1)
    similarity = pd.DataFrame(cosine_s_matrix)
    display(similarity)
    similarity.to_csv("SIMILARITY_MATRIX.csv", index=False)
    return cosine_s_matrix


def tf_idf_cosine_similarity_single_doc(a):
    cosine_s_matrix = cosine_similarity(a)
    similarity = pd.DataFrame(cosine_s_matrix)
    display(similarity)
    similarity.to_csv("SIMILARITY_MATRIX.csv", index=False)
    return cosine_s_matrix
# --------------------------------------------------Hierarchical Clustering---------------------------------------------


def ahc(similarity_matrix):
    var = linkage(similarity_matrix, 'ward')
    display(pd.DataFrame(var, columns=['Document\Cluster 1', 'Document\Cluster 2', 'Distance', 'Cluster Size'], dtype='object'))
    return var


def ahc_dendrogram(linkage_matrix):
    plt.figure(figsize=(8, 3))
    plt.title('test')
    plt.xlabel('data point')
    plt.ylabel('Distance')
    dendrogram(linkage_matrix)
    plt.axhline(y=1.0, c='k', ls='--', lw=0.5)
    plt.savefig("dendogram.pdf")


def use_of_fcluster(linkage_matrix, initial_doc):
    print("f_cluster")
    max_dist = 1.2
    cluster_labels = fcluster(linkage_matrix, max_dist, criterion='distance')
    cluster_labels = pd.DataFrame(cluster_labels, columns=['ClusterLabel'])
    x = pd.concat([initial_doc, cluster_labels], axis=1)
    display(x)

# ----------------------------------LDA TOPIC MODELING ALGORITHM--------------------------------------------------------

def lda_model(text):
    text = sentence_tokenization(text)
    text = normalizace_textu(text)
    y = [string for string in text if string != '']

    for j in range(len(y)):
        print(y[j])

    y = lemmatize(" ".join(y))
    print("after lemma", y)
    tf_idf_vector = TfidfVectorizer(min_df=0., max_df=1., norm="l2", use_idf=True, smooth_idf=True)
    tf_idf_arr = tf_idf_vector.fit_transform(y).toarray()
    vocab_tf_idf = tf_idf_vector.get_feature_names_out()
    print(vocab_tf_idf)

    print("LDA")
    lda = LatentDirichletAllocation(n_components=3, max_iter=1000, random_state=0)
    dt_matrix = lda.fit_transform(tf_idf_arr)
    features = pd.DataFrame(dt_matrix, columns=['T1', 'T2', 'T3'])
    display(features)
    topic_matrix = lda.components_

    for topic_weights in topic_matrix:
        topic = [(token, weight) for token, weight in zip(vocab_tf_idf, topic_weights)]
        topic = sorted(topic, key=lambda x: -x[1])
        topic = [item for item in topic if item[1] > 0.8]
        print(topic)
        print()
# ---------------------------------WORD2VEC-----------------------------------------------------------------------------

# def w2v(text):
#
#     tokens = [nltk.tokenize(b) for b in text]
#     print(tokens)
#     feature_size = 100
#     window_context = 30
#     min_word_count = 1
#     sample = 1e-3
#
#     w2v_model = word2vec.Word2Vec(tokens, vector_size=feature_size, window=window_context, min_count=min_word_count, sample=sample, epochs=50)
#
#     similar_words = {search_term: [item[0] for item in w2v_model.wv.most_similar([search_term], topn=5)] for search_term in ['structure']}
#     print(similar_words)





if __name__ == "__main__":
    a = 0
    x = read_docx('Sample_text_1.docx', paragraph=True)
    # y = read_docx('Sample text 2.docx', paragraph=True)
    # x = multiple_text('Sample_text_1.docx', 'Sample text 2.docx')
    # ("Enter how many most used words do you want to see.")
    # BOW_M, CV = BOW(x, int(input()))
    # # tf_idf_t(BOW_M,CV)
    c, a, DF1 = tf_idf_v(x)
    # DF1.to_csv("TFIDFV.csv", index=False)
    # b, DF2 = tf_idf_v(y)
    # DF2.to_csv("TFIDFV2.csv", index=False)
    # print(type(a))
    # print(len(a[1]))
    # print(len(b[1]))
    # tf_idf_cosine_similarity(a, b)

